"""Document loaders for PDF, DOCX, and TXT files."""

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import BinaryIO


@dataclass
class LoadedDocument:
    text: str
    filename: str
    file_type: str
    char_count: int
    metadata: dict


def _clean_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_pdf(file_obj: BinaryIO, filename: str) -> LoadedDocument:
    from PyPDF2 import PdfReader

    reader = PdfReader(file_obj)
    pages: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        pages.append(page_text)

    text = _clean_text("\n\n".join(pages))
    return LoadedDocument(
        text=text,
        filename=filename,
        file_type="pdf",
        char_count=len(text),
        metadata={"page_count": len(reader.pages)},
    )


def load_docx(file_obj: BinaryIO, filename: str) -> LoadedDocument:
    from docx import Document

    doc = Document(file_obj)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = _clean_text("\n\n".join(paragraphs))
    return LoadedDocument(
        text=text,
        filename=filename,
        file_type="docx",
        char_count=len(text),
        metadata={"paragraph_count": len(paragraphs)},
    )


def load_txt(file_obj: BinaryIO, filename: str) -> LoadedDocument:
    raw = file_obj.read()
    if isinstance(raw, bytes):
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                text = raw.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = raw.decode("utf-8", errors="replace")
    else:
        text = raw

    text = _clean_text(text)
    return LoadedDocument(
        text=text,
        filename=filename,
        file_type="txt",
        char_count=len(text),
        metadata={"encoding": "utf-8"},
    )


def load_from_bytes(content: bytes, filename: str) -> LoadedDocument:
    suffix = Path(filename).suffix.lower()
    buffer = BytesIO(content)

    if suffix == ".pdf":
        return load_pdf(buffer, filename)
    if suffix == ".docx":
        return load_docx(buffer, filename)
    if suffix == ".txt":
        return load_txt(buffer, filename)

    raise ValueError(f"Unsupported file type: {suffix}. Allowed: .pdf, .docx, .txt")


def load_from_path(path: Path) -> LoadedDocument:
    with open(path, "rb") as f:
        return load_from_bytes(f.read(), path.name)


def load_from_text(text: str, source_name: str) -> LoadedDocument:
    cleaned = _clean_text(text)
    return LoadedDocument(
        text=cleaned,
        filename=source_name,
        file_type="text",
        char_count=len(cleaned),
        metadata={"source": "raw_text"},
    )
